from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "docs" / "Rovvy_Six_Month_26_Week_Execution_Plan.docx"

TEAL = "0F766E"
NAVY = "0F172A"
SURFACE = "1E293B"
MINT = "CCFBF1"
PALE = "F8FAFC"
SLATE = "64748B"
MUTED = "94A3B8"
WHITE = "FFFFFF"
GREEN = "22C55E"
AMBER = "F59E0B"
RED = "EF4444"

weeks = [
    ("Foundation and baseline", "Confirm the product baseline, active routes, current quality, and six-month decision rules.",
     "Inventory solo, group, and join-a-group journeys; freeze the active navigation model; identify duplicate and dead UI.",
     "Run lint, type-check, tests, accessibility scan, dependency review, and production configuration audit.",
     "Create a route-by-route UI inventory using the existing Rovvy palette, typography, spacing, and component patterns.",
     "Recruit 12 design partners: four solo travelers, four existing groups, and four group-seeking travelers.",
     "Approved baseline report; prioritized backlog; named owners; analytics event dictionary; zero ambiguity about MVP scope.",
     "100% active routes inventoried; 12 interviews scheduled; baseline funnel and reliability numbers recorded.",
     "Scope inflation before evidence", "No feature enters the six-month critical path without a user problem, owner, metric, and acceptance test."),
    ("Design system and navigation", "Make Rovvy feel like one product without changing its established brand identity.",
     "Finalize Explore, Live, Trips, Connect, and Profile information architecture with contextual access to bookings and money.",
     "Consolidate semantic design tokens and shared primitives; retain meaningful map, chart, transport, and status colors.",
     "Standardize page headers, cards, buttons, forms, tabs, dialogs, empty states, and mobile navigation.",
     "Run five first-click tests on the new navigation and compare completion time with the baseline.",
     "Shared UI primitives adopted by the highest-traffic screens; documented exceptions for immersive Live mode.",
     "90% first-click success on core destinations; no contrast failures in primary flows; no brand-color drift.",
     "A cosmetic redesign that ignores behavior", "Accept UI work only when it reduces confusion, inconsistency, or task time."),
    ("Analytics and consent", "Build trustworthy measurement before optimizing acquisition or engagement.",
     "Define activation separately for solo, existing-group, and join-a-group users; add consent-aware tracking.",
     "Implement event schemas, server-side outcome events, privacy controls, retention limits, and dashboard-ready identifiers.",
     "Design privacy explanations and location-sharing consent states in plain language.",
     "Test event naming and consent comprehension with six users; verify that declining optional analytics preserves core use.",
     "Validated event stream for signup, profile completion, trip creation, invite, join request, first plan, and expense settlement.",
     ">=95% event delivery in staging; duplicate rate <2%; all sensitive events pass privacy review.",
     "Collecting precise location without sufficient purpose", "Minimize data, make sharing time-bound, and log consent changes."),
    ("Onboarding by travel mode", "Give each traveler a short, relevant path to first value.",
     "Introduce three onboarding intents: travel solo, plan with my group, or find compatible people/groups.",
     "Persist onboarding intent and progressive profile completion without blocking exploration.",
     "Create concise onboarding screens, skip paths, progress feedback, and accessible validation.",
     "Usability-test all three paths with at least three participants per intent.",
     "Users land in the correct home context with one clear next action and recoverable incomplete profiles.",
     ">=70% onboarding completion; median completion under three minutes; <10% validation abandonment.",
     "Asking for trust-heavy details too early", "Collect only what is necessary for first value; defer sensitive fields."),
    ("Solo discovery MVP", "Make Rovvy immediately useful before a traveler has created or joined a group.",
     "Strengthen destination discovery, activities, events, saved places, and lightweight day planning.",
     "Normalize explore data, improve empty/error handling, cache safe responses, and add observability for provider failures.",
     "Unify discovery cards, filters, map/list transitions, save actions, and itinerary prompts.",
     "Run task tests for finding, comparing, saving, and revisiting an activity in an unfamiliar city.",
     "A solo traveler can discover three relevant options, save one, and turn it into a plan without a group.",
     ">=40% of activated solo users save an item; median discovery-to-save under five minutes.",
     "Generic recommendations with no situational relevance", "Start with reliable filters and context before heavy personalization."),
    ("Live companion reliability", "Turn Live into a dependable in-trip assistance surface rather than a visual demo.",
     "Prioritize nearby context, route clarity, saved places, meet points, and quick recovery when connectivity is weak.",
     "Harden geocoding, route providers, map loading, permissions, offline states, and battery-conscious location updates.",
     "Preserve immersive Live styling while standardizing controls, panels, labels, and error feedback.",
     "Field-test walking and transit scenarios in two environments with permission denial and network loss cases.",
     "Live launches quickly, explains its state, and always offers a safe fallback when data is unavailable.",
     "Map usable p75 <3 seconds on supported networks; route failure recovery >=95%; zero blocked exits.",
     "False confidence from stale or incomplete location data", "Show freshness, accuracy, provider state, and manual alternatives."),
    ("Trip creation and planning", "Make starting a trip simple enough for the organizer and understandable to every member.",
     "Streamline trip creation, dates, destination, visibility, group association, and first planning action.",
     "Validate state transitions, permissions, date constraints, and API errors; add idempotency where duplicate creation is possible.",
     "Create a focused trip overview with next actions, progress, people, bookings, plans, and money.",
     "Observe four organizers creating a real or realistic trip and inviting their first collaborator.",
     "A new organizer creates a trip and reaches a collaborative state in one uninterrupted session.",
     ">=60% trip-creation completion; median under four minutes; duplicate-trip rate below 1%.",
     "Overloading the initial form", "Use progressive setup and move optional planning detail into the trip workspace."),
    ("Invites and group activation", "Make the invitation loop fast, clear, and measurable.",
     "Improve invite links/codes, pending states, resend behavior, role clarity, and non-member previews.",
     "Secure invite tokens, rate-limit abuse, instrument delivery/acceptance, and handle expired or revoked invitations.",
     "Design organizer feedback and recipient landing pages with the trip value visible before signup.",
     "Test inviter and recipient halves separately; include a traveler without an existing Rovvy account.",
     "Organizers know exactly who has been invited, joined, declined, or needs follow-up.",
     ">=35% invite acceptance within seven days; <5% invalid-link failures; supportable recovery for every state.",
     "Invite friction hidden behind authentication", "Preserve context through signup and return users to the exact invitation."),
    ("Collaborative decisions", "Reduce chat chaos by making group choices explicit and fair.",
     "Connect polls to trip decisions for destination, date, activity, budget, and meeting choices.",
     "Harden voting rules, guest voting, auto-close behavior, notifications, and concurrency.",
     "Show decision status, participation, deadlines, and the selected outcome inside the trip.",
     "Run a five-person planning simulation and measure time to reach a decision.",
     "A group creates, votes, closes, and converts a decision into a plan without leaving Rovvy.",
     ">=50% of active multi-member trips create one decision; >=65% invited members participate.",
     "Polls become another disconnected feature", "Every poll must attach to a trip context and produce a next action."),
    ("Shared itinerary and ownership", "Create a single source of truth for what the group is doing and who owns each action.",
     "Unify scheduled items, notes, meet points, reservations, assignments, and change visibility.",
     "Define ordering, timezone behavior, edit permissions, optimistic updates, and conflict resolution.",
     "Design day views, agenda summaries, ownership labels, and mobile quick actions.",
     "Test a three-day itinerary with edits from organizer and member roles across time zones.",
     "Members can answer what, when, where, and who without reading an external chat history.",
     ">=50% of activated group trips add three itinerary items; edit failure rate <2%.",
     "Timezone and concurrent-edit confusion", "Store canonical times, display local context, and preserve change history."),
    ("Bookings hub", "Make travel reservations visible and actionable without pretending to be a full OTA.",
     "Consolidate flights, hotels, buses, and routes under Trips > Bookings with manual and integration-ready records.",
     "Normalize booking data, protect confirmation details, and add provider-independent status models.",
     "Create booking cards, traveler assignments, missing-detail prompts, and chronological grouping.",
     "Test adding and locating a booking during a simulated airport disruption.",
     "The group can see critical booking facts and ownership from one trip context.",
     ">=30% of active trips add a booking; critical-detail retrieval under 20 seconds in testing.",
     "Expanding prematurely into transaction fulfillment", "Stay coordination-first until demand and partner economics are proven."),
    ("Money and settlement", "Make shared travel spending understandable, fair, and safe.",
     "Unify expense capture, splits, balances, settlement suggestions, currencies, and activity history.",
     "Audit decimal precision, currency conversion, authorization, concurrency, migrations, and reconciliation tests.",
     "Use clear payer/participant language, balance summaries, edit consequences, and accessible charts only where useful.",
     "Run realistic scenarios: equal split, custom split, multi-currency, refund, member removal, and correction.",
     "Every balance is explainable from source expenses and changes are traceable.",
     "100% deterministic scenario reconciliation; zero unauthorized edits; calculation coverage on critical paths.",
     "Financial trust loss from a single unexplained balance", "Prioritize auditability and clarity above clever automation."),
    ("Connect profiles and compatibility", "Help travelers evaluate fit without reducing people to superficial matching scores.",
     "Define opt-in travel preferences, pace, budget style, interests, languages, accessibility needs, and boundaries.",
     "Separate public, connection-only, and private fields; add moderation-ready profile controls.",
     "Design respectful compatibility explanations and profile completeness cues without dark patterns.",
     "Interview group-seeking travelers about what creates confidence and what feels invasive.",
     "A traveler can express enough context to make a safe, informed connection decision.",
     ">=60% of Connect users complete core preference fields; privacy comprehension >=80% in testing.",
     "Sensitive profiling or discriminatory filters", "Use optional, user-controlled signals and review every filter for harm."),
    ("Join-a-group workflow", "Turn discovery into a transparent request, review, and acceptance process.",
     "Build browse, preview, request-to-join, organizer review, questions, approve/decline, and onboarding states.",
     "Harden trip-only versus group membership semantics, permissions, invite codes, and notification delivery.",
     "Make expectations, visibility, cost, dates, participant count, and decision status clear before requesting.",
     "Run end-to-end simulations with applicant, organizer, accepted member, and declined applicant roles.",
     "No participant is surprised by what joining grants, exposes, costs, or requires.",
     ">=20% qualified preview-to-request rate; median organizer response under 48 hours in pilot.",
     "Treating access approval as sufficient trust", "Pair joining with safety education, boundaries, and reporting controls."),
    ("Trust, safety, and moderation", "Make safety a product system, not a policy page added after growth.",
     "Add report, block, leave, remove, emergency guidance, community standards, and incident escalation paths.",
     "Implement evidence retention rules, moderation queues, access logging, rate limits, and abuse monitoring.",
     "Place safety controls where incidents happen; use neutral language and avoid exposing reporters.",
     "Conduct tabletop exercises for harassment, impersonation, unsafe meetup, payment dispute, and compromised account.",
     "A documented severity model, response SLA, owner rotation, and auditable case lifecycle exist before public matching.",
     "100% critical scenarios have a runbook; block/report accessible within two interactions.",
     "Launching social discovery without operational capacity", "Limit pilot size to what the moderation team can safely support."),
    ("Notifications and communication", "Deliver timely coordination without becoming another noisy group chat.",
     "Prioritize invitations, decision deadlines, itinerary changes, meet points, safety events, and expense actions.",
     "Create preference controls, batching, deduplication, delivery status, retries, and channel fallbacks.",
     "Design actionable notifications with clear source, trip context, urgency, and deep link.",
     "Test interruption tolerance and notification preferences across all three traveler modes.",
     "Critical events reach the right person while low-value activity is summarized or muted.",
     ">=90% critical delivery success; duplicate notifications <1%; opt-out reasons captured.",
     "Notification fatigue destroys trust", "Default to fewer high-value notifications and let users control channels."),
    ("AI companion guardrails", "Use AI where it reduces planning effort while keeping facts, consent, and control visible.",
     "Focus AI on trip summaries, option comparison, planning prompts, and contextual assistance—not autonomous commitments.",
     "Add grounding, provider citations where available, sensitive-data rules, rate controls, eval cases, and fallback behavior.",
     "Clearly label generated content, allow editing, and show when information may be stale.",
     "Evaluate 100 representative prompts across solo, group, live, booking, money, and safety contexts.",
     "AI suggestions are reviewable, reversible, and never silently modify shared plans or financial data.",
     ">=85% useful-response rating in curated evals; 0 critical safety-policy failures.",
     "Hallucinated travel or safety information", "Ground high-impact claims and route emergencies to authoritative sources."),
    ("Closed alpha", "Move from feature completion to observed, supported travel behavior.",
     "Enroll 30-50 participants across the three segments and define weekly missions.",
     "Create release channels, feature flags, diagnostics, backup procedures, and fast rollback.",
     "Add in-product feedback, bug-report context, and a lightweight alpha status surface.",
     "Run founder-led onboarding and two structured feedback sessions during the week.",
     "Every alpha participant has a segment, intended journey, support channel, and consented research plan.",
     ">=60% complete their segment activation; p1 bugs acknowledged within four hours.",
     "Confusing compliments with evidence", "Judge value by repeated behavior, task completion, and willingness to invite others."),
    ("Alpha learning sprint", "Convert observed friction into prioritized fixes instead of adding breadth.",
     "Repair the top activation, collaboration, and trust breakdowns found in alpha.",
     "Triage defects by severity and frequency; add regression tests before closing critical issues.",
     "Polish only the journeys demonstrated to matter; remove or defer distracting surfaces.",
     "Conduct ten short playback interviews using real session evidence and support tickets.",
     "A written continue/change/stop decision exists for each major product hypothesis.",
     "Activation improves >=15% from alpha baseline; critical open defect count trends downward.",
     "Founder preference overrides user evidence", "Use the agreed scorecard and document exceptions explicitly."),
    ("Partner and supply pilot", "Test whether local and travel partners improve utility and distribution.",
     "Define partner value for activities, host communities, accommodations, universities, and travel clubs.",
     "Create controlled partner metadata, attribution, lead tracking, and content-quality controls.",
     "Build a lightweight partner profile or offer treatment consistent with Rovvy—not advertising clutter.",
     "Interview ten potential partners and pilot with two that fit the target traveler journeys.",
     "A repeatable partner proposition, onboarding checklist, and measurable traveler benefit are documented.",
     "Two pilots live; partner-sourced activation measurable; no pay-to-rank ambiguity.",
     "Partnership work distracts from product-market fit", "Choose partners only when they improve core traveler outcomes."),
    ("Pricing and willingness to pay", "Test monetization hypotheses without damaging early trust or network growth.",
     "Compare free coordination, organizer Pro, premium Live/AI assistance, and partner revenue hypotheses.",
     "Validate subscription states, entitlements, Stripe webhooks, cancellation, retries, and sandbox reconciliation.",
     "Create honest upgrade messaging tied to outcomes; never lock safety or basic joining behind payment.",
     "Run pricing interviews and landing-page tests; do not infer willingness from compliments.",
     "A pricing recommendation includes target buyer, value metric, free boundary, price range, and evidence strength.",
     ">=15 pricing conversations; measurable intent signal; zero entitlement leakage in tests.",
     "Monetizing before retention exists", "Treat pricing as discovery unless repeat usage supports a paid beta."),
    ("Growth loops and referrals", "Grow through successful travel collaboration rather than paid acquisition alone.",
     "Improve invite, shared trip preview, post-trip recap, referral, and community ambassador loops.",
     "Add attribution that respects consent and avoids exposing private trip information.",
     "Design share artifacts with safe defaults, clear audience controls, and Rovvy brand consistency.",
     "Test three messages with activated users and measure actual sharing, not stated preference.",
     "One primary organic loop is selected with a known trigger, receiver value, and conversion event.",
     "Invite rate >=0.8 per activated organizer; referral acceptance improves week over week.",
     "Viral mechanics leak private travel data", "Default shared artifacts to minimum information and explicit audience selection."),
    ("Performance, security, and accessibility", "Prepare the product for a wider beta with measurable quality gates.",
     "Close high-severity usability and accessibility defects across all active routes.",
     "Run dependency, auth, authorization, secret, rate-limit, database, backup, performance, and recovery reviews.",
     "Verify keyboard, focus, contrast, responsive layout, reduced motion, and screen-reader behavior.",
     "Include users with accessibility needs and low-connectivity conditions in validation.",
     "No known critical security issue; recovery tested; active journeys meet the beta quality threshold.",
     "Zero open p0/p1 security defects; WCAG AA on core flows; API p95 targets documented and met.",
     "Shipping known high-impact defects to meet a date", "Quality gates outrank calendar pressure for safety and data integrity."),
    ("Beta readiness and operations", "Make launch supportable, observable, and reversible.",
     "Finalize beta scope, eligibility, onboarding, help content, community standards, and feedback channels.",
     "Complete monitoring, alerts, incident response, backups, rollback, data export/deletion, and capacity tests.",
     "Polish release notes, status messaging, support entry points, and known-limitation communication.",
     "Run a full launch rehearsal with product, engineering, safety, support, and founder roles.",
     "A signed go/no-go checklist, duty roster, escalation tree, and daily launch dashboard are ready.",
     "All launch blockers closed; alert tests pass; support SLA and ownership confirmed.",
     "Operational ambiguity during incidents", "Every alert has an owner, threshold, runbook, and communication path."),
    ("Controlled beta launch", "Release to a bounded audience and learn without losing operational control.",
     "Open access in cohorts across all three segments; preserve feature flags and eligibility controls.",
     "Monitor errors, latency, abuse, notification delivery, data consistency, and provider health continuously.",
     "Resolve launch-critical comprehension issues and maintain consistent release communication.",
     "Hold daily user playback and support reviews; contact users who abandon critical steps.",
     "The beta operates within reliability and safety capacity while producing comparable cohort data.",
     "Activation >=55%; week-one retained core action >=30%; no unresolved critical incident.",
     "Scaling signups faster than support capacity", "Increase cohorts only after the previous cohort clears quality gates."),
    ("Retention, fundraising, and next-quarter decision", "Turn six months of behavior into a defensible retention model and evidence-based investment plan.",
     "Identify the strongest repeat-value loop by segment; synthesize product proof, market wedge, business model, and next milestones.",
     "Build cohort views and document architecture health, security posture, scalability constraints, technical debt, and hiring needs.",
     "Refine lifecycle experiences and prepare coherent investor demo flows using actual product evidence and Rovvy identity.",
     "Interview retained and churned users, then conduct advisory reviews and three investor-style rehearsals.",
     "Board-ready six-month review, north-star behavior, 12-month roadmap, hiring plan, use of funds, and go/change/stop decisions.",
     "Week-two repeat action >=20% in the target cohort; validated traction narrative and next-quarter objectives approved.",
     "Fundraising story or aggregate metrics outrun segment evidence", "Lead with observed behavior and review all three traveler cohorts independently."),
]

assert len(weeks) == 26

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar"); fldChar.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instruction
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t"); txt.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (fldChar, instr, sep, txt, end): run._r.append(el)

def set_font(run, size=10.5, color=NAVY, bold=False, italic=False, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold; run.italic = italic

def add_p(doc, text="", size=10.5, color=NAVY, bold=False, italic=False, after=5, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None: p.alignment = align
    set_font(p.add_run(text), size, color, bold, italic)
    return p

def add_bullet(doc, label, detail, color=NAVY):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.06
    set_font(p.add_run(label + ": "), 9.7, color, True)
    set_font(p.add_run(detail), 9.7, NAVY)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(7 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, 17 if level == 1 else 12, TEAL if level == 1 else SURFACE, True, name="Outfit")
    return p

def page_label(doc, kicker, title, subtitle=None):
    add_p(doc, kicker.upper(), 8.5, TEAL, True, after=3)
    add_p(doc, title, 22, NAVY, True, after=4)
    if subtitle: add_p(doc, subtitle, 10.5, SLATE, False, after=11)

def add_callout(doc, label, text, fill=MINT, accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False; table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0); cell.width = Inches(6.5); shade(cell, fill); set_cell_margins(cell, 120, 150, 120, 150)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(label.upper() + "  "), 8.5, accent, True)
    set_font(p.add_run(text), 10.2, NAVY, True)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def page_break(doc):
    doc.add_page_break()

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.72); sec.bottom_margin = Inches(0.68)
sec.left_margin = Inches(1); sec.right_margin = Inches(1)
sec.header_distance = Inches(0.35); sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(NAVY)
normal.paragraph_format.space_after = Pt(5); normal.paragraph_format.line_spacing = 1.08
for n, size in ((1, 17), (2, 12), (3, 10.5)):
    st = styles[f"Heading {n}"]
    st.font.name = "Outfit"; st.font.size = Pt(size); st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(TEAL if n == 1 else SURFACE)
    st.paragraph_format.keep_with_next = True
    st.paragraph_format.space_before = Pt(8); st.paragraph_format.space_after = Pt(4)
lb = styles["List Bullet"]
lb.font.name = "Aptos"; lb.font.size = Pt(9.7)
lb.paragraph_format.left_indent = Inches(0.26); lb.paragraph_format.first_line_indent = Inches(-0.16)

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("ROVVY  /  SIX-MONTH EXECUTION PLAN"), 8, MUTED, True)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("CONFIDENTIAL  •  "), 8, MUTED, True)
add_field(footer, "PAGE")
set_font(footer.add_run("  OF  "), 8, MUTED, True)
add_field(footer, "NUMPAGES")

# Page 1
add_p(doc, "ROVVY", 12, TEAL, True, after=80, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, "Six-Month Product, Market\nand Company Execution Plan", 30, NAVY, True, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, "26 weeks from product coherence to controlled beta and investor readiness", 14, SLATE, False, after=45, align=WD_ALIGN_PARAGRAPH.CENTER)
add_callout(doc, "Product mission", "A trusted travel companion for people traveling solo, coordinating with an existing group, or safely finding and joining compatible travelers.")
add_p(doc, "Planning horizon: 10 August 2026 – 7 February 2027", 11, NAVY, True, before=36, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, "Prepared as an operating roadmap for founders, product, engineering, design, growth, safety, and investors", 9.5, MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
page_break(doc)

# Page 2
page_label(doc, "How to use this plan", "Purpose, assumptions, and decision discipline")
add_callout(doc, "Core assumption", "Rovvy wins by coordinating the human reality of travel—not by becoming another undifferentiated booking marketplace.")
for label, detail in [
    ("Audience", "Founders, product and engineering leads, design, growth, operations, trust and safety, advisors, and prospective investors."),
    ("Cadence", "Every week ends with a review of shipped outcomes, user evidence, quality metrics, risks, and the next gate."),
    ("Segments", "Metrics must remain separated for solo travelers, existing groups, and travelers seeking to join a new group."),
    ("Scope rule", "No work enters the critical path without a defined problem, owner, user outcome, metric, and acceptance test."),
    ("Release rule", "Safety, privacy, authorization, financial integrity, and recovery gates outrank calendar commitments."),
    ("Evidence rule", "Behavioral data and observed usability carry more weight than feature count, compliments, or founder preference."),
    ("Brand rule", "Preserve Rovvy teal, navy, mint, typography, tone, and recognizable identity across every surface."),
    ("Planning rule", "This is a living operating document; revise sequence when evidence changes, but record the reason and expected impact."),
]: add_bullet(doc, label, detail)
page_break(doc)

# Page 3
page_label(doc, "Product strategy", "The three-sided traveler promise")
for title, promise, outcome in [
    ("Traveling solo", "Useful before a network exists", "Discover relevant places, build a lightweight plan, navigate confidently, save context, and get assistance during the trip."),
    ("Traveling with an existing group", "One shared operating layer", "Coordinate decisions, itinerary, people, bookings, meet points, notifications, and money without fragmented chats and spreadsheets."),
    ("Joining new travelers or groups", "Compatibility with accountable trust", "Understand the trip and people, request access transparently, set boundaries, and retain safety controls before and after joining."),
]:
    add_heading(doc, title, 2); add_p(doc, promise, 10.5, TEAL, True, after=3); add_p(doc, outcome, 10.2, NAVY, after=8)
add_callout(doc, "Shared platform advantage", "The same identity, trip context, map, plan, communication, safety, and payment records can support all three modes while each mode keeps a distinct activation journey.")
page_break(doc)

# Page 4
page_label(doc, "Positioning", "Market wedge and strategic boundaries")
add_heading(doc, "Positioning statement", 2)
add_p(doc, "For travelers who need confidence and coordination before and during a trip, Rovvy is a travel companion that connects discovery, live context, group decisions, shared plans, people, and money. Unlike booking-first tools or generic group chat, Rovvy organizes the trip relationship itself.", 11, NAVY, after=12)
for label, detail in [
    ("Initial wedge", "Small leisure groups and socially open travelers planning multi-day trips, where coordination pain and trust questions are both visible."),
    ("Do now", "Discovery, trip workspaces, collaboration, bookings visibility, expenses, live assistance, joining workflows, safety, and measurable onboarding."),
    ("Do later", "Deep transaction fulfillment, broad marketplace supply, complex loyalty programs, autonomous AI actions, and native apps without validated demand."),
    ("Never compromise", "Location privacy, clear membership boundaries, financial auditability, consent, safety escalation, and honest AI behavior."),
    ("Competitive frame", "Compete on integrated coordination and trust, not on inventory breadth or lowest-price search."),
    ("Investor proof", "Show segment-specific activation, collaborative usage, invite conversion, repeated trip value, safety readiness, and a credible monetization path."),
]: add_bullet(doc, label, detail)
page_break(doc)

# Page 5
page_label(doc, "Six-month outcomes", "What must be demonstrably true by Week 26")
outcomes = [
    ("Coherent product", "Every active route follows one Rovvy design system and a clear Explore / Live / Trips / Connect architecture."),
    ("Three activations", "Each traveler mode has a short, instrumented path to first value and its own conversion baseline."),
    ("Collaborative value", "Groups repeatedly use decisions, itinerary, people, bookings, and money in a shared trip context."),
    ("Trustworthy connection", "Join-a-group pilots operate with visibility controls, reporting, blocking, moderation, and response runbooks."),
    ("Reliable operations", "Core journeys meet agreed performance, accessibility, security, backup, monitoring, and incident standards."),
    ("Market evidence", "Closed alpha and controlled beta cohorts generate retention, invitation, willingness-to-pay, and partner evidence."),
    ("Investment readiness", "Rovvy has an evidence-led narrative, 12-month roadmap, hiring plan, use of funds, and explicit strategic decisions."),
]
for label, detail in outcomes: add_bullet(doc, label, detail, TEAL)
add_callout(doc, "End-state test", "A founder should be able to demonstrate a complete solo journey, a complete existing-group journey, and a safe join-a-group journey using real product behavior—not presentation-only prototypes.")
page_break(doc)

# Page 6
page_label(doc, "Company scorecard", "Metrics that govern weekly decisions")
score = [
    ("Activation", "Intent selected → profile sufficient → first segment-specific value action completed."),
    ("Collaboration", "Trips with two or more participating members and at least one shared decision or plan action."),
    ("Invite loop", "Invites per activated organizer, acceptance rate, time to acceptance, and recipient activation."),
    ("Connect quality", "Qualified previews, requests, organizer response time, acceptance, blocks/reports, and post-join health."),
    ("Retention", "Repeat core action by segment in week one and week two; next-trip return when observable."),
    ("Reliability", "Crash-free sessions, API error rate, p75/p95 latency, provider failure recovery, and notification delivery."),
    ("Trust", "Safety incidents by severity, response SLA, privacy-control use, unauthorized access attempts, and resolution quality."),
    ("Business", "Pricing interview evidence, upgrade intent, pilot partner contribution, acquisition source, and support cost."),
]
for label, detail in score: add_bullet(doc, label, detail)
add_p(doc, "All scorecard views must be filterable by traveler mode, acquisition cohort, platform, and release cohort. Aggregate growth must never hide a failing or unsafe segment.", 9.8, SLATE, italic=True, before=5)
page_break(doc)

# Page 7
page_label(doc, "Operating model", "Workstreams, ownership, and weekly rhythm")
for label, detail in [
    ("Product", "Owns problem definition, sequencing, acceptance criteria, segment metrics, and decisions to continue/change/stop."),
    ("Engineering", "Owns architecture, delivery, tests, security, performance, observability, data integrity, and reversible releases."),
    ("Design", "Owns interaction coherence, responsive behavior, accessibility, research prototypes, and the fixed Rovvy identity."),
    ("Growth and market", "Owns recruiting, interviews, messaging, referral experiments, partnerships, pricing evidence, and cohort quality."),
    ("Trust and operations", "Owns consent, moderation, escalation, support, incident response, community standards, and launch capacity."),
    ("Founder", "Owns strategic boundaries, capital allocation, hiring, investor communication, and final go/no-go decisions."),
    ("Monday", "Confirm the weekly outcome, leading metric, release scope, owner, dependencies, and kill conditions."),
    ("Wednesday", "Review working software and user evidence; remove blockers and prevent scope creep."),
    ("Friday", "Demo shipped behavior, inspect scorecard and incidents, record learning, and approve the next gate."),
]: add_bullet(doc, label, detail)
page_break(doc)

# Page 8
page_label(doc, "Roadmap overview", "Six phases across 26 weeks")
phases = [
    ("1. Foundation", "Weeks 1–4", "Baseline, design system, analytics, segmented onboarding"),
    ("2. Core utility", "Weeks 5–8", "Solo discovery, Live, trip creation, invitations"),
    ("3. Group operating layer", "Weeks 9–12", "Decisions, itinerary, bookings, money"),
    ("4. Connection and trust", "Weeks 13–17", "Profiles, joining, safety, notifications, AI guardrails"),
    ("5. Evidence and business", "Weeks 18–22", "Alpha, iteration, partners, pricing, growth loops"),
    ("6. Beta and investment", "Weeks 23–26", "Hardening, launch operations, beta, retention, fundraising"),
]
table = doc.add_table(rows=1, cols=3)
table.autofit = False
widths = [1.5, 1.05, 3.95]
for i, w in enumerate(widths): table.columns[i].width = Inches(w)
hdr = table.rows[0]; set_repeat_table_header(hdr)
for i, text in enumerate(("Phase", "Timing", "Primary outcome")):
    cell = hdr.cells[i]; shade(cell, NAVY); set_cell_margins(cell)
    set_font(cell.paragraphs[0].add_run(text), 9, WHITE, True)
for phase, timing, outcome in phases:
    cells = table.add_row().cells
    for i, text in enumerate((phase, timing, outcome)):
        cells[i].width = Inches(widths[i]); shade(cells[i], PALE if len(table.rows) % 2 == 0 else WHITE); set_cell_margins(cells[i])
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_font(cells[i].paragraphs[0].add_run(text), 9.2, NAVY, i == 0)
add_callout(doc, "Critical sequence", "Connection discovery does not scale before privacy and moderation; monetization does not lead before retention; beta does not open before reliability and incident readiness.")
page_break(doc)

# Pages 9-60: 26 weeks x 2 pages
for idx, w in enumerate(weeks, start=1):
    title, objective, product, engineering, design, validation, deliverable, metric, risk, guardrail = w
    page_label(doc, f"Week {idx:02d} / 26 • Execution", title, objective)
    add_callout(doc, "Weekly outcome", objective)
    add_heading(doc, "Workstream commitments", 2)
    add_bullet(doc, "Product", product)
    add_bullet(doc, "Engineering and data", engineering)
    add_bullet(doc, "Design and accessibility", design)
    add_bullet(doc, "Research and market", validation)
    add_heading(doc, "Execution sequence", 2)
    add_bullet(doc, "Monday", "Confirm user problem, acceptance criteria, leading metric, responsible owner, dependencies, and release boundary.")
    add_bullet(doc, "Tuesday–Wednesday", "Build the smallest end-to-end behavior, review working software early, and instrument the outcome—not only clicks.")
    add_bullet(doc, "Thursday", "Run accessibility, permission, error, empty, loading, mobile, and recovery checks; conduct the scheduled user validation.")
    add_bullet(doc, "Friday", "Release behind the appropriate control, inspect evidence, document decisions, and prepare the next gate.")
    page_break(doc)

    page_label(doc, f"Week {idx:02d} / 26 • Control sheet", "Deliverables, evidence, risks, and exit gate")
    add_heading(doc, "Definition of done", 2)
    add_bullet(doc, "Primary deliverable", deliverable, TEAL)
    add_bullet(doc, "Evidence threshold", metric, TEAL)
    add_bullet(doc, "Quality checks", "No newly introduced lint, type, test, accessibility, authorization, privacy, or responsive-layout regression in the affected journey.")
    add_bullet(doc, "Documentation", "Update decision log, event dictionary, acceptance evidence, known limitations, support notes, and the next dependency.")
    add_heading(doc, "Risk control", 2)
    add_bullet(doc, "Most likely failure", risk, RED)
    add_bullet(doc, "Guardrail", guardrail, TEAL)
    add_bullet(doc, "Rollback trigger", "Any critical safety, privacy, financial-integrity, authorization, data-loss, or unrecoverable journey failure.")
    add_heading(doc, "Friday exit questions", 2)
    for question in [
        "Did the intended traveler complete the outcome in working software?",
        "What behavior changed, and how strong is the evidence by segment?",
        "What failed or surprised us, including support and safety signals?",
        "What will we continue, change, stop, or defer next week?",
        "Is the next gate safe, measurable, owned, and reversible?",
    ]:
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3); set_font(p.add_run(question), 9.6, NAVY)
    if idx != len(weeks): page_break(doc)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
from zipfile import ZipFile
with ZipFile(OUT) as package:
    xml = package.read("word/document.xml")
    assert package.testzip() is None
    assert xml.count(b'w:type="page"') == 59
print(f"{OUT} (60 planned pages; 59 explicit page breaks; package integrity OK)")
